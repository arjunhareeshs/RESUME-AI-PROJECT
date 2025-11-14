from typing import List, Dict

from docx import Document


def _paragraph_blocks(doc: Document) -> List[Dict]:
    """
    Convert paragraphs into pseudo-blocks. DOCX does not expose coordinates,
    so we provide order-based positions and mark them as single-column content.
    """
    blocks = []
    order = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        blocks.append({
            "text": text,
            "x0": 0,
            "y0": float(order),
            "x1": 100,
            "y1": float(order + 1),
            "column": 0,
            "type": "paragraph",
        })
        order += 1
    return blocks


def _table_blocks(doc: Document, start_order: int) -> List[Dict]:
    """
    Extract tables cell by cell to avoid losing information such as multi-column layouts.
    We treat each cell as an individual block and note its row/column position.
    """
    blocks = []
    order = start_order
    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                if not text:
                    continue
                blocks.append({
                    "text": text,
                    "x0": float(col_idx * 100),
                    "y0": float(order),
                    "x1": float((col_idx + 1) * 100),
                    "y1": float(order + 1),
                    "column": col_idx,
                    "type": "table_cell",
                    "table_row": row_idx,
                    "table_col": col_idx,
                })
                order += 1
    return blocks


def extract(file_path: str) -> List[Dict]:
    """
    Extract textual content from a DOCX file, capturing paragraphs and tables.
    Column metadata is inferred where possible (tables); otherwise default to single column.
    """
    doc = Document(file_path)
    paragraph_blocks = _paragraph_blocks(doc)
    table_blocks = _table_blocks(doc, start_order=len(paragraph_blocks))

    blocks = paragraph_blocks + table_blocks
    # Sort by synthetic y coordinate to preserve reading order
    blocks.sort(key=lambda b: (b.get("column", 0), b["y0"]))
    return blocks
